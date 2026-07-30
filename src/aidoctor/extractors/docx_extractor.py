"""DOCX extraction that keeps headings attached to their prose.

python-docx exposes a flat paragraph list, so the structure a reader sees is
lost unless you rebuild it. Here a heading opens a new section and the
paragraphs beneath it accumulate into that section, which means a chunk cites
"Installation" rather than "paragraph 47" — and the chunker never splits a
heading away from the text it introduces.

Tables are read too. Skipping them is a common silent gap: requirements and
limits are exactly the sort of thing that lives in a table.
"""

from __future__ import annotations

from pathlib import Path

from aidoctor.extractors.base import ExtractionError
from aidoctor.models.document import Document, Section, SourceType


def _is_heading(style_name: str) -> bool:
    lowered = (style_name or "").lower()
    return lowered.startswith("heading") or lowered in {"title", "subtitle"}


class DocxExtractor:
    source_type = SourceType.DOCX
    extensions = (".docx",)

    def extract(self, path: Path, doc_id: str) -> Document:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ExtractionError("python-docx is required to read .docx files") from exc

        try:
            source = DocxDocument(str(path))
        except Exception as exc:
            raise ExtractionError(f"Could not open {path.name}: {exc}") from exc

        sections: list[Section] = []
        heading = "body"
        buffer: list[str] = []

        def flush() -> None:
            if buffer:
                sections.append(Section(text="\n".join(buffer).strip(), label=heading, ordinal=len(sections)))
                buffer.clear()

        for paragraph in source.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if _is_heading(getattr(paragraph.style, "name", "")):
                flush()
                heading = text
                buffer.append(text)
            else:
                buffer.append(text)
        flush()

        for index, table in enumerate(source.tables, start=1):
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()) for row in table.rows
            ]
            rows = [r for r in rows if r]
            if rows:
                sections.append(Section(text="\n".join(rows), label=f"table {index}", ordinal=len(sections)))

        if not sections:
            raise ExtractionError(f"{path.name} contained no readable text or tables")
        return Document(
            doc_id=doc_id,
            filename=path.name,
            source_type=self.source_type,
            sections=sections,
            metadata={"tables": str(len(source.tables))},
        )
